import io
from datetime import date, timedelta, datetime
import requests

import streamlit as st
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import locale
import json

# ----------------- REGRAS FIXAS DO SEU CALENDÁRIO -----------------


# ----------------- CARREGAR CALENDÁRIO DO JSON -----------------
with open("calendario.json", "r", encoding="utf-8") as f:
    calendario = json.load(f)

INICIO = datetime.strptime(calendario["inicio"], "%Y-%m-%d").date()
FIM = datetime.strptime(calendario["fim"], "%Y-%m-%d").date()

FERIADOS = {datetime.strptime(d, "%Y-%m-%d").date() for d in calendario["feriados"]}

RECESSOS = [
    (datetime.strptime(r[0], "%Y-%m-%d").date(),
     datetime.strptime(r[1], "%Y-%m-%d").date())
    for r in calendario["recessos"]
]

DIAS_NAO_LETIVOS = {datetime.strptime(d, "%Y-%m-%d").date() for d in calendario["dias_nao_letivos"]}

ETAPA1 = tuple(datetime.strptime(d, "%Y-%m-%d").date() for d in calendario["etapas"]["etapa1"])
ETAPA2 = tuple(datetime.strptime(d, "%Y-%m-%d").date() for d in calendario["etapas"]["etapa2"])

AVALIACOES_FIXAS = {
    datetime.strptime(d, "%Y-%m-%d").date(): txt
    for d, txt in calendario["avaliacoes"].items()
}

LOGO_URL = "https://raw.githubusercontent.com/ledicefreitas/CRONOGRAMA/refs/heads/main/logo%20expoente.png"

# ----------------- LÓGICA DE GERAÇÃO -----------------
def gerar_datas(inicio, fim, dias_semana_aulas, feriados, recessos, dias_nao_letivos, total_aulas, compensacoes):
    datas = []
    aulas_geradas = 0
    atual = inicio
    comp_dict = {orig: weekday_comp for orig, weekday_comp in compensacoes}

    while atual <= fim and aulas_geradas < total_aulas:
        weekday = atual.weekday()
        if atual in comp_dict:
            comp_weekday = comp_dict[atual]
            if comp_weekday in dias_semana_aulas:
                qtd_aulas = dias_semana_aulas[comp_weekday]
                for _ in range(qtd_aulas):
                    if aulas_geradas < total_aulas and atual not in datas:
                        datas.append(atual)
                        aulas_geradas += 1
            atual += timedelta(days=1)
            continue

        if weekday in dias_semana_aulas:
            em_recesso = any(r[0] <= atual <= r[1] for r in recessos)
            if atual not in feriados and not em_recesso and atual not in dias_nao_letivos:
                qtd_aulas = dias_semana_aulas[weekday]
                for _ in range(qtd_aulas):
                    if aulas_geradas < total_aulas:
                        datas.append(atual)
                        aulas_geradas += 1
        atual += timedelta(days=1)

    datas.sort()
    return datas

def definir_bordas(celula, tamanho=4, cor="000000"):
    tc = celula._tc
    tcPr = tc.get_or_add_tcPr()
    tblBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(tamanho))
        border.set(qn('w:color'), cor)
        tblBorders.append(border)
    tcPr.append(tblBorders)

_CM_TO_TWIPS = 567

def fix_table_grid(table, widths_cm):
    tbl = table._tbl
    tblGrid = OxmlElement('w:tblGrid')
    for w in widths_cm:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(int(w * _CM_TO_TWIPS)))
        tblGrid.append(gridCol)
    tbl.insert(0, tblGrid)

def adicionar_tabela_etapa(doc, titulo_etapa, periodo, datas_etapa, inicio_index, footer_text=""):
    table = doc.add_table(rows=1, cols=5)
    table.autofit = False
    table.allow_autofit = False

    col_widths = [2.5, 1.5, 12, 2.5, 2.5]
    fix_table_grid(table, col_widths)

    for i, w in enumerate(col_widths):
        for cell in table.columns[i].cells:
            cell.width = Cm(w)

    # Cabeçalho mesclado azul
    hdr = table.rows[0].cells
    hdr[0].merge(hdr[-1])
    p = hdr[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{titulo_etapa} - De {periodo[0].strftime('%d/%m/%y')} a {periodo[1].strftime('%d/%m/%y')}")
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(11)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), "0A1F44")
    hdr[0]._tc.get_or_add_tcPr().append(shading)
    definir_bordas(hdr[0])

    # Linha de títulos
    headers = ["DATA", "AULA", "CONTEÚDO", "MATERIAL DE APOIO", "AVALIAÇÃO"]
    row_hdr = table.add_row().cells
    for i, h in enumerate(headers):
        row_hdr[i].text = h
        for cell in table.columns[i].cells:
            definir_bordas(cell)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(10)
                    r.font.bold = True

    # Linha de mês
    ultimo_mes = None
    for idx, d in enumerate(datas_etapa, start=inicio_index):
        if ultimo_mes != d.month:
            meses_pt = [
                "Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho",
                "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro"
            ]
            mes_texto = f"{meses_pt[d.month - 1]}/{d.year}".upper()
            month_row = table.add_row().cells
            month_row[0].merge(month_row[-1])
            p_mes = month_row[0].paragraphs[0]
            p_mes.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_mes = p_mes.add_run(mes_texto)
            run_mes.font.bold = True
            run_mes.font.name = "Arial"
            run_mes.font.size = Pt(10)
            shading = OxmlElement('w:shd')
            month_row[0]._tc.get_or_add_tcPr().append(shading)
            definir_bordas(month_row[0])
            ultimo_mes = d.month

        # Linha normal da aula
        row_cells = table.add_row().cells
        row_cells[0].text = d.strftime("%d/%m/%Y")
        row_cells[1].text = str(idx)
        row_cells[2].text = AVALIACOES_FIXAS.get(d, "")
        row_cells[3].text = ""
        row_cells[4].text = ""
        for i, cell in enumerate(row_cells):
            definir_bordas(cell)
            cell.width = Cm(col_widths[i])
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(10)

    # Rodapé
    if footer_text:
        footer_row = table.add_row().cells
        footer_row[0].merge(footer_row[-1])
        tc = footer_row[0]._tc
        for p in tc.xpath(".//w:p"):
            tc.remove(p)
        for linha in footer_text.split("\n"):
            p = footer_row[0].add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(linha)
            run.font.name = "Arial"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 0)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), "D9D9D9")
        footer_row[0]._tc.get_or_add_tcPr().append(shading)
        definir_bordas(footer_row[0])

    return inicio_index + len(datas_etapa)


def gerar_docx(disciplina, curso, professor, turma, total_aulas, dias_semana_dict, compensacoes):
    datas_aulas = gerar_datas(
        INICIO, FIM,
        dias_semana_dict, FERIADOS, RECESSOS, DIAS_NAO_LETIVOS,
        total_aulas, compensacoes
    )

    # separa por etapas usando as constantes ETAPA1 e ETAPA2
    datas_etapa1 = [d for d in datas_aulas if ETAPA1[0] <= d <= ETAPA1[1]]
    datas_etapa2 = [d for d in datas_aulas if ETAPA2[0] <= d <= ETAPA2[1]]

    rodape_etapa1 = (
        "Obs: Na 1ª etapa serão trabalhadas 02 práticas de formação: \n"
        "1ª prática – deverá ser aplicada até o dia 05/09/25 \n"
        "2ª prática – deverá ser aplicada até o dia 03/10/25 \n"
        "As datas das práticas devem constar no cronograma de aulas."
    )
    rodape_etapa2 = (
        "Obs: Na 2ª etapa serão trabalhadas 02 práticas de formação: \n"
        "1ª prática – deverá ser aplicada até o dia 21/11/25 \n"
        "2ª prática – deverá ser aplicada até o dia 06/02/26 \n"
        "As datas das práticas devem constar no cronograma de aulas."
    )

    def criar_doc(etapa_nome, periodo, datas, inicio_index, rodape):
        doc = Document()
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width

        # Cabeçalho (mesma estrutura que tu já tinha)
        table_header = doc.add_table(rows=1, cols=2)
        table_header.autofit = False
        table_header.allow_autofit = False
        fix_table_grid(table_header, [3.5, 17.5])
        for i, w in enumerate([3.5, 17.5]):
            for cell in table_header.columns[i].cells:
                cell.width = Cm(w)

        # Logo
        try:
            response = requests.get(LOGO_URL)
            if response.status_code == 200:
                logo_bytes = io.BytesIO(response.content)
                cell_logo = table_header.rows[0].cells[0]
                cell_logo.text = ""
                p_logo = cell_logo.add_paragraph()
                p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_logo = p_logo.add_run()
                run_logo.add_picture(logo_bytes, width=Cm(2.5))
        except Exception:
            pass

        # Texto da direita (titulo e info)
        cell_info = table_header.rows[0].cells[1]
        # mantém os ajustes de margem que já usavas
        try:
            cell_info.top_margin = Pt(0)
            cell_info.bottom_margin = Pt(0)
            cell_info.left_margin = Pt(0)
            cell_info.right_margin = Pt(0)
        except Exception:
            # alguns ambientes de python-docx podem não expor essas props, ok pular
            pass

        p1 = cell_info.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_before = Pt(0)
        p1.paragraph_format.space_after = Pt(0)

        for linha in [
            "COLÉGIO EXPOENTE",
            f"CURSO TÉCNICO EM {curso}".upper(),
            f"CRONOGRAMA {etapa_nome} - 2º PERÍODO - 2025"
        ]:
            run = p1.add_run(linha + "\n")
            run.font.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(12)

        p2 = cell_info.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(0)

        ultimas = [
            f"DISCIPLINA: {disciplina}".upper(),
            f"Professor(a): {professor}".upper(),
            f"TURMA: {turma}                                               CARGA HORARIA: {total_aulas}h/a".upper()
        ]

        for i, linha in enumerate(ultimas):
            texto = linha if i == len(ultimas)-1 else linha + "\n"
            run = p2.add_run(texto)
            run.font.bold = False
            run.font.name = "Arial"
            run.font.size = Pt(12)

        for row in table_header.rows:
            for cell in row.cells:
                definir_bordas(cell)

        doc.add_paragraph("\n")

        # aqui usamos a função já existente que cria a tabela e numeracão
        adicionar_tabela_etapa(doc, etapa_nome, periodo, datas, inicio_index, footer_text=rodape)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    # calcula o início da numeração da etapa 2: começa logo após o fim da etapa 1
    inicio_etapa2 = 1 + len(datas_etapa1)

    buffer1 = criar_doc("1ª ETAPA", ETAPA1, datas_etapa1, inicio_index=1, rodape=rodape_etapa1)
    buffer2 = criar_doc("2ª ETAPA", ETAPA2, datas_etapa2, inicio_index=inicio_etapa2, rodape=rodape_etapa2)

    return buffer1, buffer2, len(datas_etapa1), len(datas_etapa2)


# ----------------- UI (Streamlit) -----------------
st.set_page_config(page_title="Gerador de Cronograma", page_icon="📅", layout="centered")
st.title("📅 Gerador Modelo de Cronograma ")
st.caption("Preencha os dados, clique em Gerar e baixe o .docx já com as \ndatas preenchidas. Fácil, rápido e sem drama 😉")

# ---- FORMULÁRIO PRINCIPAL (somente campos obrigatórios) ----
#with st.form("form"):
#    col1, col2 = st.columns(2)
#    with col1:
#        disciplina = st.text_input("Disciplina*", "")
#        curso = st.text_input("Curso*", "Vendas")
#        professor = st.text_input("Professor(a)* (Nome Completo)", "")
#    with col2:
#        turma = st.text_input("Turma*", "")
#        total_aulas = st.number_input("Número total de aulas*", min_value=1, step=1, value=30)
#    st.form_submit_button("Formulário preenchido → continue abaixo")

# ---- CAMPOS PRINCIPAIS ----
col1, col2 = st.columns(2)
with col1:
    disciplina = st.text_input("Disciplina*", "")
    curso = st.text_input("Curso*", "Vendas")
    professor = st.text_input("Professor(a)* (Nome Completo)", "")
with col2:
    turma = st.text_input("Turma*", "")
    total_aulas = st.number_input("Número total de aulas*", min_value=1, step=1, value=30)

# ---- DIAS DA SEMANA (fora do form) ----
st.markdown("**Selecione os dias da semana e quantidade de aulas**")
dias_semana_dict = {}
dias = ["Segunda", "Terça", "Quarta", "Quinta", "Sexta", "Sábado", "Domingo"]
for i, dia in enumerate(dias):
    marcar = st.checkbox(dia, key=f"dia_{i}")
    if marcar:
        qtd = st.number_input(f"Aulas na {dia}", min_value=1, step=1, key=f"aulas_{i}")
        dias_semana_dict[i] = qtd

# ---- COMPENSAÇÕES ----
st.markdown("**Compensações** no formato `dd/mm/aaaa->n` (n = 0 seg ... 6 dom). Ex.: `10/10/2025->2`")

comps_txt = st.text_input(" ")

# ----------------- FUNÇÃO AUXILIAR -----------------
def parse_compensacoes(txt: str):
    res = []
    if not txt.strip():
        return res
    for part in txt.split(","):
        if "->" not in part:
            continue
        data_str, wd_str = part.split("->")
        data_dt = datetime.strptime(data_str.strip(), "%d/%m/%Y").date()
        wd = int(wd_str.strip())
        if wd < 0 or wd > 6:
            raise ValueError("Compensação com weekday inválido (0..6).")
        res.append((data_dt, wd))
    return res

compensacoes = parse_compensacoes(comps_txt)

# ---- BOTÃO FINAL PARA GERAR CRONOGRAMA ----
if st.button("Gerar cronograma"):
    docx_etapa1, docx_etapa2, total_etapa1, total_etapa2 = gerar_docx(
        disciplina=disciplina.strip(),
        curso=curso.strip(),
        professor=professor.strip(),
        turma=turma.strip(),
        total_aulas=int(total_aulas),
        dias_semana_dict=dias_semana_dict,
        compensacoes=compensacoes
    )
    st.session_state["docx_etapa1"] = docx_etapa1
    st.session_state["docx_etapa2"] = docx_etapa2
    st.session_state["total_etapa1"] = total_etapa1
    st.session_state["total_etapa2"] = total_etapa2
    st.success(
        f"Cronogramas gerados! \n"
        f"**Total de aulas Etapa 1**: {total_etapa1} aulas. \n"
        f"**Total de aulas Etapa 2**: {total_etapa2} aulas."
    )


# mostra os botões sempre que já tiver arquivos na session
if "docx_etapa1" in st.session_state and "docx_etapa2" in st.session_state:
    st.download_button(
        "⬇️ Baixar 1ª Etapa",
        data=st.session_state["docx_etapa1"],
        file_name=f"{professor.strip().replace(' ', '_')} - CRONO - {disciplina.strip().replace(' ', '_')}-T{turma.strip().replace(' ', '_')}.docx".upper(),
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    st.download_button(
        "⬇️ Baixar 2ª Etapa",
        data=st.session_state["docx_etapa2"],
        file_name=f"{professor.strip().replace(' ', '_')} - CRONO - {disciplina.strip().replace(' ', '_')}-T{turma.strip().replace(' ', '_')}.docx".upper(),
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


    
